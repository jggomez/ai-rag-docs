import io
import logging
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


class DocxResponseGenerator:
    """
    Generates a formal response letter Word (.docx) file using python-docx.

    The document is composed of two parts:
      1. Response letter page(s) – the AI-generated reply.
      2. References page        – a structured summary of every context
                                  document (received chunks + resolved sent
                                  responses) that was used to produce the reply.
    """

    def generate_response_docx(
        self,
        response_text: str,
        metadata: Dict[str, str],
        similar_chunks: Optional[List[dict]] = None,
        sent_texts: Optional[Dict[str, dict]] = None,
    ) -> bytes:
        """
        Generates a Word document with the provided response text, metadata, and
        an optional references page listing all context documents used.

        Args:
            response_text:  The AI-generated response text body.
            metadata:       Engineering metadata (contract_number, sender, …).
            similar_chunks: Reranked received-document chunks used as RAG context.
            sent_texts:     Dict mapping {id_borrador: {texto, filename}} of resolved
                            sent responses used as RAG context.

        Returns:
            Word file contents as bytes.
        """
        doc = Document()

        # Set standard margins (1 inch)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Configure default Normal style font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)  # Text dark color

        # ── Letter Section ──
        subject = metadata.get("subject", "")
        if subject:
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.space_after = Pt(12)
            run_lbl = p_ref.add_run("Ref: ")
            run_lbl.bold = True
            p_ref.add_run(subject)

        for paragraph in response_text.split("\n"):
            clean = paragraph.strip()
            if clean:
                p = doc.add_paragraph(clean)
                p.paragraph_format.space_after = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # ── References Section ──
        if similar_chunks:
            doc.add_page_break()

            # Page Title
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_after = Pt(4)
            run_title = p_title.add_run("Documentos de Contexto Utilizados")
            run_title.bold = True
            run_title.size = Pt(14)
            run_title.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)  # Dark Blue

            # Caption
            p_caption = doc.add_paragraph(
                "La siguiente sección presenta los documentos de referencia que el sistema "
                "utilizó como contexto para generar la respuesta anterior."
            )
            p_caption.paragraph_format.space_after = Pt(24)
            p_caption.runs[0].font.size = Pt(9)
            p_caption.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Sub-section 1: Received documents
            p_sub1 = doc.add_paragraph()
            p_sub1.paragraph_format.space_after = Pt(12)
            run_sub1 = p_sub1.add_run("Documentos Recibidos (contexto de recuperación)")
            run_sub1.bold = True
            run_sub1.size = Pt(11)
            run_sub1.font.color.rgb = RGBColor(0x2c, 0x6f, 0xad)  # Mid Blue

            # Table for received chunks
            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            # Set column widths
            col_widths = [Inches(0.4), Inches(3.8), Inches(1.6), Inches(0.7)]

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "#"
            hdr_cells[1].text = "Asunto"
            hdr_cells[2].text = "Documento"
            hdr_cells[3].text = "Relevancia"

            # Style header row
            for idx, cell in enumerate(hdr_cells):
                cell.width = col_widths[idx]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

            # Insert data rows
            for idx, chunk in enumerate(similar_chunks, 1):
                row_cells = table.add_row().cells
                rerank_score = chunk.get("rerank_score")
                score_str = f"{rerank_score:.3f}" if rerank_score is not None else "—"

                row_cells[0].text = str(idx)
                row_cells[1].text = chunk.get("asunto", "—")
                row_cells[2].text = chunk.get("nombre_archivo", "—")
                row_cells[3].text = score_str

                # Apply widths and basic formatting
                for cell_idx, cell in enumerate(row_cells):
                    cell.width = col_widths[cell_idx]
                    p = cell.paragraphs[0]
                    if len(p.runs) > 0:
                        p.runs[0].font.size = Pt(8)
                    if cell_idx in (0, 3):
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add spaces
            doc.add_paragraph().paragraph_format.space_after = Pt(12)

            # Sub-section 2: Sent responses
            seen_draft_ids = set()
            ordered_sent = []
            for chunk in similar_chunks:
                draft_id = chunk.get("id_borrador")
                if draft_id and draft_id in (sent_texts or {}) and draft_id not in seen_draft_ids:
                    seen_draft_ids.add(draft_id)
                    data = sent_texts[draft_id]
                    ordered_sent.append({
                        "filename": data.get("filename", draft_id),
                        "texto": data.get("texto", ""),
                        "asunto": chunk.get("asunto", "—"),
                    })

            if ordered_sent:
                p_sub2 = doc.add_paragraph()
                p_sub2.paragraph_format.space_after = Pt(12)
                run_sub2 = p_sub2.add_run("Respuestas Enviadas (contexto de generación)")
                run_sub2.bold = True
                run_sub2.size = Pt(11)
                run_sub2.font.color.rgb = RGBColor(0x2c, 0x6f, 0xad)

                for entry in ordered_sent:
                    # Header for sent document
                    p_sent_hdr = doc.add_paragraph()
                    p_sent_hdr.paragraph_format.space_after = Pt(2)
                    run_sent_lbl = p_sent_hdr.add_run(
                        f"Documento: {entry['filename']}  |  Asunto: {entry['asunto']}"
                    )
                    run_sent_lbl.bold = True
                    run_sent_lbl.font.size = Pt(9)
                    run_sent_lbl.font.color.rgb = RGBColor(0x1e, 0x6e, 0x42)  # Green

                    # Sent document excerpt
                    excerpt = entry["texto"][:600].strip()
                    if len(entry["texto"]) > 600:
                        excerpt += "…"

                    p_sent_body = doc.add_paragraph(excerpt)
                    p_sent_body.paragraph_format.space_after = Pt(12)
                    p_sent_body.runs[0].font.size = Pt(9)

        # Write doc to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()

        logger.info(
            f"Generated Word docx response ({len(docx_bytes)} bytes, "
            f"{len(similar_chunks or [])} context docs referenced)."
        )
        return docx_bytes
